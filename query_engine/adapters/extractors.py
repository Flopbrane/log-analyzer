"""外部文書ファイルから検索用TextDocumentを作成する抽出器。"""
from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path
from typing import IO, TYPE_CHECKING, Protocol

from query_engine.adapters.documents import TextDocument, from_text, normalize_text

TEXT_EXTENSIONS: frozenset[str] = frozenset({".txt", ".md", ".rst", ".csv", ".json", ".jsonl", ".log", ".py", ".js", ".ts"})
HTML_EXTENSIONS: frozenset[str] = frozenset({".html", ".htm", ".xhtml"})
PDF_EXTENSIONS: frozenset[str] = frozenset({".pdf"})
DOCX_EXTENSIONS: frozenset[str] = frozenset({".docx"})

if TYPE_CHECKING:
    from bs4 import BeautifulSoup
    from charset_normalizer.models import CharsetMatch, CharsetMatches
    from docx.document import Document as DocxDocument
    from pypdf import PdfReader


class CharsetNormalizerModule(Protocol):
    def from_bytes(self, sequences: bytes | bytearray) -> CharsetMatches: ...


class Bs4Module(Protocol):
    BeautifulSoup: type[BeautifulSoup]


class PdfModule(Protocol):
    PdfReader: type[PdfReader]


class DocxModule(Protocol):
    def Document(self, docx: str | IO[bytes] | None = None) -> DocxDocument: ...


def _import_charset_normalizer() -> CharsetNormalizerModule:
    try:
        import charset_normalizer
    except ImportError as exc:
        raise RuntimeError("charset-normalizer が必要です。pip install charset-normalizer を実行してください。") from exc
    return charset_normalizer


def _import_bs4() -> Bs4Module:
    try:
        import bs4
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 が必要です。pip install beautifulsoup4 を実行してください。") from exc
    return bs4


def _import_pypdf() -> PdfModule:
    try:
        import pypdf
    except ImportError as exc:
        raise RuntimeError("pypdf が必要です。pip install pypdf を実行してください。") from exc
    return pypdf


def _import_docx() -> DocxModule:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx が必要です。pip install python-docx を実行してください。") from exc
    return docx


def _extract_pdf(path: Path, **metadata: object) -> TextDocument:
    pypdf_module: PdfModule = _import_pypdf()
    reader: PdfReader = pypdf_module.PdfReader(str(path))
    parts: list[str] = [page.extract_text() or "" for page in reader.pages]
    return from_text("\n".join(parts), title=path.stem, source=str(path), **metadata)


def _extract_docx(path: Path, **metadata: object) -> TextDocument:
    docx_module: DocxModule = _import_docx()
    document: DocxDocument = docx_module.Document(str(path))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    return from_text("\n".join(parts), title=path.stem, source=str(path), **metadata)


def extract_text_file(path: str | Path, **metadata: object) -> TextDocument:
    """拡張子に応じて文書を読み込み、TextDocumentへ変換する。"""
    file_path: Path = Path(path)
    suffix: str = file_path.suffix.casefold()
    if suffix in TEXT_EXTENSIONS:
        return _extract_plain_text(file_path, **metadata)
    if suffix in HTML_EXTENSIONS:
        return _extract_html(file_path, **metadata)
    if suffix in PDF_EXTENSIONS:
        return _extract_pdf(file_path, **metadata)
    if suffix in DOCX_EXTENSIONS:
        return _extract_docx(file_path, **metadata)
    raise ValueError(f"未対応のファイル形式です: {suffix or file_path.name}")


def _extract_plain_text(path: Path, **metadata: object) -> TextDocument:
    charset_normalizer: CharsetNormalizerModule = _import_charset_normalizer()
    raw: bytes = path.read_bytes()
    detected: CharsetMatch | None = charset_normalizer.from_bytes(raw).best()
    text: str = "" if detected is None else str(detected)
    return from_text(text, title=path.stem, source=str(path), **metadata)


def _extract_html(path: Path, **metadata: object) -> TextDocument:
    charset_normalizer: CharsetNormalizerModule = _import_charset_normalizer()
    raw: bytes = path.read_bytes()
    detected: CharsetMatch | None = charset_normalizer.from_bytes(raw).best()
    html: str = "" if detected is None else str(detected)
    try:
        bs4_module: Bs4Module = _import_bs4()
    except RuntimeError:
        title: str
        text: str
        title, text = _extract_html_with_stdlib(html, default_title=path.stem)
        return from_text(text, title=title, source=str(path), **metadata)
    soup: BeautifulSoup = bs4_module.BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title = normalize_text(soup.title.get_text(" ")) if soup.title else path.stem
    text = normalize_text(soup.get_text(" "))
    return from_text(text, title=title or path.stem, source=str(path), **metadata)


def _extract_html_with_stdlib(html: str, *, default_title: str) -> tuple[str, str]:
    parser: _PlainTextHTMLParser = _PlainTextHTMLParser()
    parser.feed(html)
    return parser.title or default_title, normalize_text(" ".join(parser.parts))


class _PlainTextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.title: str = ""
        self._in_title = False
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag == "title":
            self._in_title = True

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        if tag == "title":
            self._in_title = False

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        text: str = normalize_text(data)
        if not text:
            return
        if self._in_title:
            self.title: str = text
        self.parts.append(text)
